from argparse import ArgumentParser, ArgumentDefaultsHelpFormatter
import argparse
from pathlib import Path
from typing import Iterator, List, Union
import docx
from simplify_docx import simplify
from itertools import chain

# TODO: write the DataLoader class then import it in data_analysis.ipynb

def main(args) -> None:
    data_path = make_data_path(args.data_dir, args.data_fn)

    book = DataLoader(data_path)

    print(book.chapters)

def by_type(paragraph: List[dict]) -> bool:
    return any(p['VALUE'] == '[w:drawing]' for p in paragraph)

def make_data_path(data_dir: str, data_fn: str) -> Path:
    data_path = Path(data_dir + data_fn).expanduser().resolve()
    assert data_path.exists()
    return data_path


def get_args() -> argparse.Namespace:
    arg_parser = ArgumentParser(formatter_class=ArgumentDefaultsHelpFormatter)
    arg_parser.add_argument('-d', '--data-dir', type=str,
                            default="data/",
                            help="Path to directory containing the input data file")
    arg_parser.add_argument('-f', "--data-fn", type=str,
                            default="D5627-Dolan.docx",
                            help=("File name with the text to summarize"))
    return arg_parser.parse_args()


class DataLoader:

    def __init__(self, data_path: Path, start_idx=155, n_chapters=5) -> None:
        assert data_path.exists()
        self.docx: dict = simplify(docx.Document(data_path))
        self.paragraphs: Iterator[Union[str, List[dict]]] = self._init_paragraphs()

        self.start_idx = start_idx
        self.n_chapters = n_chapters
        self.chapters = self._init_chapters()

    def _init_chapters(self, marker="Chapitre {} /"):
        _paragraphs: List[Union[str, List[dict]]] = list(self.paragraphs)
        chapters: Union[str, List[dict]] = []

        start, end = 0, self.start_idx
        bound = lambda i: _paragraphs.index(marker.format(i), self.start_idx)
        for chapter_idx in range(1, self.n_chapters + 1):
            chapters.append(_paragraphs[start: end])
            start = bound(chapter_idx)
            end = bound(chapter_idx + 1) if chapter_idx < self.n_chapters else None
        chapters.append(_paragraphs[start:])

        # TODO: Look into concatenating each chapter

        return chapters


    def _init_paragraphs(self) -> Iterator[Union[str, List[dict]]]:
        _docx: List[dict] = self.docx["VALUE"][0]["VALUE"]
        _paragraphs = map(lambda p: p["VALUE"], _docx)
        _paragraphs = filter(lambda p: p != "[w:sdt]", _paragraphs)
        _paragraphs = chain.from_iterable(_paragraphs)
        _paragraphs = filter(lambda p: p["TYPE"] != "CT_Empty", _paragraphs)

        return map(lambda p: p["VALUE"], _paragraphs)


if __name__ == "__main__":
    main(get_args())
